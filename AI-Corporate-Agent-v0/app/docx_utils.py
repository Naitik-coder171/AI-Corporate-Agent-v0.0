from __future__ import annotations
from typing import List, Dict, Any
from pathlib import Path
from docx import Document  # type: ignore
from docx.enum.text import WD_COLOR_INDEX  # type: ignore
from docx.shared import Pt  # type: ignore


def annotate_docx(input_path: str, issues: List[Dict[str, Any]], output_path: str) -> None:
    doc = Document(input_path)

    # Add a summary section at the top
    heading = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    run = heading.insert_paragraph_before("ADGM Review Notes:").runs[0]
    run.bold = True
    run.font.size = Pt(12)

    for idx, issue in enumerate(issues, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"[{idx}] {issue.get('section_hint', 'General')}: {issue.get('issue', '')} ")
        r.bold = True
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if issue.get("suggestion"):
            doc.add_paragraph(f"Suggestion: {issue['suggestion']}")
        if issue.get("citation"):
            doc.add_paragraph(f"Citation: {issue['citation']}")

    # Try to highlight occurrences of section hints
    for issue in issues:
        hint = (issue.get("section_hint") or "").strip()
        if not hint:
            continue
        for para in doc.paragraphs:
            if hint.lower() in para.text.lower():
                for run in para.runs:
                    if hint.lower() in run.text.lower():
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        run.bold = True

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path) 