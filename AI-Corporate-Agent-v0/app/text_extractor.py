from __future__ import annotations
from pathlib import Path
from typing import Tuple
import re

from docx import Document  # type: ignore
from PyPDF2 import PdfReader  # type: ignore


def extract_text_with_metadata(path: str) -> Tuple[str, dict]:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".docx":
        return _extract_docx(p), {"source": str(p), "type": "docx"}
    if suffix == ".pdf":
        return _extract_pdf(p), {"source": str(p), "type": "pdf"}
    # Fallback for md/txt
    text = p.read_text(encoding="utf-8", errors="ignore")
    return text, {"source": str(p), "type": suffix.lstrip('.') or 'text'}


def _extract_docx(p: Path) -> str:
    doc = Document(str(p))
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    # Include tables
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    text = "\n".join(parts)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def _extract_pdf(p: Path) -> str:
    try:
        reader = PdfReader(str(p))
        texts = []
        for page in reader.pages:
            txt = page.extract_text() or ""
            texts.append(txt)
        return "\n".join(texts)
    except Exception:
        # If PDF parsing fails, return empty to avoid crashing
        return "" 